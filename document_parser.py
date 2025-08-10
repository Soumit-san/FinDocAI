import PyPDF2
import io
import docx
import pdfplumber
from typing import Optional

class DocumentParser:
    def __init__(self):
        pass
    
    def parse_document(self, uploaded_file) -> str:
        """Parse uploaded documents and extract text content"""
        try:
            file_extension = uploaded_file.name.split('.')[-1].lower()
            
            if file_extension == 'pdf':
                return self._parse_pdf(uploaded_file)
            elif file_extension == 'txt':
                return self._parse_txt(uploaded_file)
            elif file_extension in ['docx', 'doc']:
                return self._parse_docx(uploaded_file)
            else:
                raise ValueError(f"Unsupported file format: {file_extension}")
                
        except Exception as e:
            raise Exception(f"Error parsing document: {str(e)}")
    
    def _parse_pdf(self, uploaded_file) -> str:
        """Parse PDF files with multiple extraction methods"""
        try:
            # Reset file pointer to beginning
            uploaded_file.seek(0)
            pdf_data = uploaded_file.read()
            text_content = ""
            
            # Method 1: Try pdfplumber first (better for complex layouts and tables)
            try:
                pdf_file = io.BytesIO(pdf_data)
                with pdfplumber.open(pdf_file) as pdf:
                    for page_num, page in enumerate(pdf.pages):
                        try:
                            page_text = page.extract_text()
                            if page_text:
                                text_content += page_text + "\n\n"
                                
                            # Also extract text from tables if present
                            tables = page.extract_tables()
                            for table in tables:
                                for row in table:
                                    if row:
                                        text_content += " | ".join([cell or "" for cell in row]) + "\n"
                                text_content += "\n"
                                
                        except Exception as page_error:
                            print(f"Warning: pdfplumber failed on page {page_num + 1}: {str(page_error)}")
                            continue
                            
                if text_content.strip():
                    return text_content.strip()
                    
            except Exception as plumber_error:
                print(f"pdfplumber extraction failed: {str(plumber_error)}")
            
            # Method 2: Fallback to PyPDF2 if pdfplumber fails
            try:
                pdf_file = io.BytesIO(pdf_data)
                pdf_reader = PyPDF2.PdfReader(pdf_file)
                
                for page_num in range(len(pdf_reader.pages)):
                    page = pdf_reader.pages[page_num]
                    try:
                        page_text = page.extract_text()
                        if page_text:
                            text_content += page_text + "\n\n"
                    except Exception as page_error:
                        print(f"Warning: PyPDF2 failed on page {page_num + 1}: {str(page_error)}")
                        continue
                        
                if text_content.strip():
                    return text_content.strip()
                    
            except Exception as pypdf_error:
                print(f"PyPDF2 extraction failed: {str(pypdf_error)}")
            
            # If both methods fail
            if not text_content.strip():
                # Try to get basic info about the PDF
                try:
                    pdf_file = io.BytesIO(pdf_data)
                    pdf_reader = PyPDF2.PdfReader(pdf_file)
                    num_pages = len(pdf_reader.pages)
                    
                    if num_pages == 0:
                        raise ValueError("PDF file appears to be empty or corrupted")
                    else:
                        raise ValueError(
                            f"Could not extract text from PDF with {num_pages} pages. "
                            "This might be a scanned PDF (image-based), encrypted, or have complex formatting. "
                            "Try: 1) Converting to Word/text format, 2) Using OCR software, or 3) Uploading a different PDF."
                        )
                except:
                    raise ValueError(
                        "PDF file could not be processed. It may be corrupted, encrypted, or in an unsupported format. "
                        "Please try uploading a different file."
                    )
            
            return text_content.strip()
            
        except Exception as e:
            # More specific error handling
            error_msg = str(e).lower()
            if "encrypted" in error_msg or "password" in error_msg:
                raise Exception("PDF is password protected. Please provide an unprotected PDF file.")
            elif "corrupted" in error_msg or "invalid" in error_msg:
                raise Exception("PDF file appears to be corrupted. Please try uploading a different file.")
            else:
                raise Exception(f"Error parsing PDF: {str(e)}")
    
    def _parse_txt(self, uploaded_file) -> str:
        """Parse text files"""
        try:
            # Read text file content
            text_content = uploaded_file.read().decode('utf-8')
            
            if not text_content.strip():
                raise ValueError("Text file appears to be empty")
            
            return text_content
            
        except UnicodeDecodeError:
            try:
                # Try different encoding
                uploaded_file.seek(0)
                text_content = uploaded_file.read().decode('latin-1')
                return text_content
            except Exception as e:
                raise Exception(f"Error decoding text file: {str(e)}")
                
        except Exception as e:
            raise Exception(f"Error parsing text file: {str(e)}")
    
    def _parse_docx(self, uploaded_file) -> str:
        """Parse DOCX files"""
        try:
            # Create a BytesIO object from uploaded file
            docx_file = io.BytesIO(uploaded_file.read())
            
            # Load document
            doc = docx.Document(docx_file)
            
            text_content = ""
            
            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                text_content += paragraph.text + "\n"
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text_content += cell.text + "\t"
                    text_content += "\n"
            
            if not text_content.strip():
                raise ValueError("No text content could be extracted from DOCX")
            
            return text_content
            
        except Exception as e:
            raise Exception(f"Error parsing DOCX: {str(e)}")
    
    def validate_financial_document(self, text_content: str) -> dict:
        """Validate if document appears to be financial in nature"""
        try:
            # Financial keywords to look for
            financial_keywords = [
                'revenue', 'earnings', 'profit', 'loss', 'balance sheet',
                'income statement', 'cash flow', 'quarterly', 'annual',
                'fiscal year', 'operating income', 'net income', 'ebitda',
                'assets', 'liabilities', 'equity', 'dividend', 'shares',
                'stock', 'investment', 'financial', 'sec filing', '10-k',
                '10-q', '8-k', 'earnings report', 'financial statements'
            ]
            
            text_lower = text_content.lower()
            found_keywords = []
            
            for keyword in financial_keywords:
                if keyword in text_lower:
                    found_keywords.append(keyword)
            
            confidence_score = len(found_keywords) / len(financial_keywords)
            is_financial = confidence_score > 0.1  # At least 10% of keywords found
            
            return {
                'is_financial_document': is_financial,
                'confidence_score': confidence_score,
                'found_keywords': found_keywords,
                'document_length': len(text_content),
                'word_count': len(text_content.split())
            }
            
        except Exception as e:
            return {
                'is_financial_document': False,
                'confidence_score': 0.0,
                'found_keywords': [],
                'document_length': 0,
                'word_count': 0,
                'error': str(e)
            }
    
    def extract_financial_metrics(self, text_content: str) -> dict:
        """Extract common financial metrics from document text"""
        try:
            import re
            
            metrics = {}
            
            # Common financial metric patterns
            patterns = {
                'revenue': r'revenue[s]?\s*[:\-]?\s*\$?([\d,]+\.?\d*)\s*(million|billion|thousand)?',
                'net_income': r'net\s+income[s]?\s*[:\-]?\s*\$?([\d,]+\.?\d*)\s*(million|billion|thousand)?',
                'earnings_per_share': r'earnings?\s+per\s+share[s]?\s*[:\-]?\s*\$?([\d,]+\.?\d*)',
                'total_assets': r'total\s+assets?\s*[:\-]?\s*\$?([\d,]+\.?\d*)\s*(million|billion|thousand)?',
                'market_cap': r'market\s+cap(?:italization)?\s*[:\-]?\s*\$?([\d,]+\.?\d*)\s*(million|billion|thousand)?'
            }
            
            text_lower = text_content.lower()
            
            for metric, pattern in patterns.items():
                matches = re.findall(pattern, text_lower)
                if matches:
                    # Take the first match
                    amount, unit = matches[0]
                    amount = amount.replace(',', '')
                    
                    # Convert to standard format
                    try:
                        value = float(amount)
                        if unit.lower() == 'billion':
                            value *= 1e9
                        elif unit.lower() == 'million':
                            value *= 1e6
                        elif unit.lower() == 'thousand':
                            value *= 1e3
                        
                        metrics[metric] = {
                            'value': value,
                            'formatted': f"${value:,.2f}",
                            'original_text': matches[0]
                        }
                    except ValueError:
                        continue
            
            return metrics
            
        except Exception as e:
            return {'error': str(e)}
    
    def get_document_summary(self, text_content: str) -> dict:
        """Get summary statistics about the document"""
        try:
            words = text_content.split()
            sentences = text_content.split('.')
            
            return {
                'character_count': len(text_content),
                'word_count': len(words),
                'sentence_count': len(sentences),
                'average_words_per_sentence': len(words) / len(sentences) if sentences else 0,
                'estimated_reading_time_minutes': len(words) / 200,  # Average reading speed
                'document_type': self._identify_document_type(text_content)
            }
            
        except Exception as e:
            return {'error': str(e)}
    
    def _identify_document_type(self, text_content: str) -> str:
        """Identify the type of financial document"""
        text_lower = text_content.lower()
        
        document_types = {
            '10-k': ['10-k', 'annual report', 'form 10-k'],
            '10-q': ['10-q', 'quarterly report', 'form 10-q'],
            '8-k': ['8-k', 'current report', 'form 8-k'],
            'earnings_report': ['earnings report', 'quarterly earnings', 'earnings call'],
            'press_release': ['press release', 'news release', 'announces'],
            'analyst_report': ['analyst report', 'research report', 'price target'],
            'proxy_statement': ['proxy statement', 'def 14a', 'shareholder meeting']
        }
        
        for doc_type, keywords in document_types.items():
            for keyword in keywords:
                if keyword in text_lower:
                    return doc_type.replace('_', ' ').title()
        
        return 'Financial Document'
